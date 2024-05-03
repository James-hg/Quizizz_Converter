import pandas as pd
import docx

docpath = "/Users/jameshoang/Downloads/On tap CAPI conv.docx"

fulltext = ""

try:
    doc = docx.Document(docpath)
except:
    print("File not exist")

for para in doc.paragraphs:
    fulltext += para.text
arr = ["A.", "B.", "C.", "D.", "Câu "]
for i in arr:
    fulltext = fulltext.replace(i, f"\n{i}")
fulltext = fulltext.replace("\t", "")
fulltext.strip()

fulltext = fulltext.splitlines()

answer = []

# (run.font.color.rgb == RGBColor(255, 000, 000)) or (run.underline) or (run.bold)
for para in docx.Document(docpath).paragraphs:
    for run in para.runs:
        if run.underline:
            if run.text.startswith("A"):
                temp = "1"
                answer.append(temp)

            elif run.text.startswith("B"):
                temp = "2"
                answer.append(temp)

            elif run.text.startswith("C"):
                temp = "3"
                answer.append(temp)

            elif run.text.startswith("D"):
                temp = "4"
                answer.append(temp)
# print(answer)
print(len(answer))


separator = "|"
final = f"Question Text{separator}Question Type{separator}Option 1{separator}Option 2{separator}Option 3{
    separator}Option 4{separator}Correct Answer{separator}Time in seconds{separator}Image Link\n"

count = 0
try:
    for i in fulltext:
        if i.startswith("Câu "):
            final += i + separator + "Multiple Choice" + separator
        elif i.startswith("A.") or i.startswith("B.") or i.startswith("C."):
            final += i + separator
        elif i.startswith("D."):
            final += i + separator + answer[count] + \
                separator + "900\n"
            count += 1
except:
    print("Answer error")

counter = 1
while counter <= count:
    ques = f"Câu {counter}: "
    final = final.replace(ques, '')
    ques = f"Câu {counter}. "
    final = final.replace(ques, '')
    counter += 1

for char in ['A.', 'B.', 'C.', 'D.']:
    final = final.replace(char, '')

rows = final.strip().split('\n')
data = [row.split('|') for row in rows]
df = pd.DataFrame(data)
try:
    df.to_excel('output.xlsx', index=False, header=False)
except:
    print("Excel error")
